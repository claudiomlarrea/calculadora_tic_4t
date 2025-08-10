
import pandas as pd
from docx import Document
from io import BytesIO
from typing import Dict, Tuple, Any, Optional

# =========================
# Configuración y etiquetas
# =========================
ETIQUETAS_BINARIAS = {
    "IP_III_04": {1: "Sí", 2: "No", 9: "Ns/Nc"},  # Uso de computadora (individuos)
    "IP_III_05": {1: "Sí", 2: "No", 9: "Ns/Nc"},  # Uso de celular (individuos)
    "IP_III_06": {1: "Sí", 2: "No", 9: "Ns/Nc"},  # Uso de internet (individuos)
    "IH_II_01": {1: "Sí", 2: "No", 9: "Ns/Nc"},  # Hogar con computadora
    "IH_II_02": {1: "Sí", 2: "No", 9: "Ns/Nc"},  # Hogar con acceso a internet
    "CH04": {1: "Varón", 2: "Mujer"},            # Sexo
}

ETIQUETAS_EDUC = {
    1: "Sin instrucción",
    2: "Primaria incompleta",
    3: "Primaria completa",
    4: "Secundaria incompleta",
    5: "Secundaria completa",
    6: "Superior incompleto",
    7: "Superior completo",
    9: "Ns/Nc",
}

CANDIDATAS_EDAD = ["CH06", "EDAD", "edad"]
CANDIDATAS_SEXO = ["CH04", "SEXO"]
CANDIDATAS_EDUC = ["NIVEL_ED", "NIVEL_EDUC", "NIVEL_EDUCATIVO"]
CANDIDATA_INGRESO = ["ITF", "ITF_HOGAR", "ingreso_total_hogar"]
CANDIDATAS_PESO = ["PONDERA", "PONDIIO", "PONDIH", "PESO", "FACTOR", "FACTOR_EXP"]

def _primera(df: pd.DataFrame, cols) -> Optional[str]:
    for c in cols:
        if c in df.columns:
            return c
    return None

def _peso(df: pd.DataFrame) -> Optional[pd.Series]:
    c = _primera(df, CANDIDATAS_PESO)
    return df[c] if c else None

def _edad_grupos(df: pd.DataFrame) -> pd.Series:
    c = _primera(df, CANDIDATAS_EDAD)
    if c and pd.api.types.is_numeric_dtype(df[c]):
        e = df[c].clip(lower=0, upper=100)
        bins = [-0.1, 17, 29, 44, 64, 150]
        labels = ["0-17", "18-29", "30-44", "45-64", "65+"]
        return pd.cut(e, bins=bins, labels=labels)
    # Fallback si no hay edad
    return pd.qcut(range(len(df)), 5, labels=["0-17", "18-29", "30-44", "45-64", "65+"])

def _quintiles_ingreso(df: pd.DataFrame) -> Optional[pd.Series]:
    c = _primera(df, CANDIDATA_INGRESO)
    if not c or not pd.api.types.is_numeric_dtype(df[c]):
        return None
    try:
        q = pd.qcut(df[c].rank(method="first"), 5, labels=["Q1 (más bajo)", "Q2", "Q3", "Q4", "Q5 (más alto)"])
        q.name = "Quintil ingreso"
        return q
    except Exception:
        return None

def _map_nominal(nombre: str, s: pd.Series) -> pd.Series:
    if nombre in ETIQUETAS_BINARIAS:
        return s.map(ETIQUETAS_BINARIAS[nombre]).fillna(s)
    if nombre in CANDIDATAS_EDUC:
        if pd.api.types.is_numeric_dtype(s):
            return s.map(ETIQUETAS_EDUC).fillna(s)
    return s

def _value_counts_w(s: pd.Series, w: Optional[pd.Series]) -> pd.DataFrame:
    if w is not None and len(w) == len(s):
        tmp = pd.DataFrame({"x": s, "w": w}).dropna(subset=["x"])
        out = tmp.groupby("x", as_index=False)["w"].sum().rename(columns={"x": s.name, "w": "Cantidad"})
        out["Cantidad"] = out["Cantidad"].round(0).astype(int)
        return out.sort_values(s.name)
    vc = s.value_counts(dropna=False).sort_index().reset_index()
    vc.columns = [s.name, "Cantidad"]
    return vc

def _prop_w(flag: pd.Series, w: Optional[pd.Series]) -> float:
    if w is not None and len(w) == len(flag):
        df = pd.DataFrame({"f": flag.astype(float), "w": w}).dropna(subset=["f"])
        num = (df["f"] * df["w"]).sum()
        den = df["w"].sum()
        return float(num/den) if den > 0 else 0.0
    return float(flag.mean()) if len(flag) else 0.0

def _tabla_prop_por(g: pd.Series, flag: pd.Series, w: Optional[pd.Series]) -> pd.DataFrame:
    data = pd.DataFrame({"g": g, "f": flag.astype(int)}).dropna(subset=["g"])
    if w is not None and len(w) == len(g):
        data["w"] = w
        agg = data.groupby("g").apply(lambda d: (d["f"]*d["w"]).sum() / d["w"].sum()).reset_index()
    else:
        agg = data.groupby("g")["f"].mean().reset_index()
    agg.columns = [g.name or "Grupo", "Porcentaje"]
    agg["Porcentaje"] = (agg["Porcentaje"] * 100).round(2)
    return agg

# =====================================
# Núcleo: cálculos y narrativa extendida
# =====================================
def generar_analisis_tic_ampliado(df_merged: pd.DataFrame) -> Tuple[Dict[str, pd.DataFrame], pd.DataFrame, Dict[str, Any]]:
    df = df_merged.copy()

    # Validación mínima
    faltantes = [v for v in ["IP_III_04", "IP_III_06"] if v not in df.columns]
    if faltantes:
        raise ValueError(f"Faltan columnas clave: {faltantes}. Revisá el instructivo del trimestre.")

    # Indicadores de exclusión
    df["excluido_binario"] = ((df["IP_III_04"] == 2) & (df["IP_III_06"] == 2)).astype(int)
    df["exclusion_ordinal"] = df.apply(
        lambda x: 2 if x["IP_III_04"] == 2 and x["IP_III_06"] == 2
        else 1 if x["IP_III_04"] == 2 or x["IP_III_06"] == 2
        else 0, axis=1
    )

    # Etiquetas
    for c in ["IP_III_04","IP_III_05","IP_III_06","IH_II_01","IH_II_02","CH04"]:
        if c in df.columns:
            df[c] = _map_nominal(c, df[c])

    edad = _edad_grupos(df); edad.name = "Edad"
    sexo_col = _primera(df, CANDIDATAS_SEXO)
    educ_col = _primera(df, CANDIDATAS_EDUC)
    quintil = _quintiles_ingreso(df)
    w = _peso(df)

    tablas: Dict[str, pd.DataFrame] = {}

    # Distribuciones básicas
    for col, titulo in [
        ("IP_III_04","Uso de computadora"),
        ("IP_III_05","Uso de celular"),
        ("IP_III_06","Uso de internet"),
        ("IH_II_01","Hogar con computadora"),
        ("IH_II_02","Hogar con acceso a internet"),
    ]:
        if col in df.columns:
            dist = _value_counts_w(df[col], w)
            dist[col] = _map_nominal(col, dist[col])
            total = dist["Cantidad"].sum()
            dist["%"] = (dist["Cantidad"] / total * 100).round(2)
            tablas[titulo] = dist

    # Exclusión ordinal
    excl_ordinal = _value_counts_w(df["exclusion_ordinal"], w)
    excl_ordinal["Nivel"] = excl_ordinal["exclusion_ordinal"].map({0:"Sin exclusión",1:"Exclusión parcial",2:"Exclusión total"})
    excl_ordinal = excl_ordinal[["Nivel","Cantidad"]]
    excl_ordinal["%"] = (excl_ordinal["Cantidad"]/excl_ordinal["Cantidad"].sum()*100).round(2)
    tablas["Exclusión – Nivel ordinal"] = excl_ordinal

    # Brechas por edad, sexo, educación
    tablas["Brecha por edad (exclusión)"] = _tabla_prop_por(edad, df["excluido_binario"], w)

    if sexo_col:
        sexolab = _map_nominal(sexo_col, df[sexo_col]).rename("Sexo")
        tablas["Brecha por sexo (exclusión)"] = _tabla_prop_por(sexolab, df["excluido_binario"], w)

    if educ_col:
        educlab = _map_nominal(educ_col, df[educ_col]).rename("Nivel educativo")
        tablas["Brecha por educación (exclusión)"] = _tabla_prop_por(educlab, df["excluido_binario"], w)

    if ("IH_II_02" in df.columns) and (quintil is not None):
        acc_int = (df["IH_II_02"] == "Sí") if df["IH_II_02"].dtype == object else (df["IH_II_02"] == 1)
        acc_quintil = _tabla_prop_por(quintil, acc_int.rename("acceso"), w)
        acc_quintil.rename(columns={"Porcentaje":"% hogares con internet"}, inplace=True)
        tablas["Acceso a internet por quintil"] = acc_quintil

    # Segmentos críticos (edad x educación x sexo) – top 5 tasas de exclusión
    segs = pd.DataFrame({"Edad": edad, "ex": df["excluido_binario"]})
    if educ_col: segs["Educación"] = _map_nominal(educ_col, df[educ_col])
    if sexo_col: segs["Sexo"] = _map_nominal(sexo_col, df[sexo_col])
    if w is not None: segs["w"] = w
    keys = [c for c in ["Edad","Educación","Sexo"] if c in segs.columns]
    if keys:
        if "w" in segs.columns:
            top = segs.dropna(subset=keys).groupby(keys).apply(lambda d: (d["ex"]*d.get("w",1)).sum()/d.get("w",1).sum()).reset_index()
        else:
            top = segs.dropna(subset=keys).groupby(keys)["ex"].mean().reset_index()
        top.columns = keys + ["Tasa exclusión (%)"]
        top["Tasa exclusión (%)"] = (top["Tasa exclusión (%)"]*100).round(2)
        top = top.sort_values("Tasa exclusión (%)", ascending=False).head(5).reset_index(drop=True)
        tablas["Top 5 segmentos con mayor exclusión"] = top

    # Resumen cuantitativo para narrativa
    resumen: Dict[str, Any] = {}
    resumen["exclusion_total_pct"] = round(_prop_w(df["excluido_binario"], w)*100, 2)
    if "IH_II_01" in df.columns:
        pc_flag = (df["IH_II_01"] == "Sí") if df["IH_II_01"].dtype == object else (df["IH_II_01"] == 1)
        resumen["hogar_con_pc_pct"] = round(_prop_w(pc_flag, w)*100, 2)
    if "IH_II_02" in df.columns:
        net_flag = (df["IH_II_02"] == "Sí") if df["IH_II_02"].dtype == object else (df["IH_II_02"] == 1)
        resumen["hogar_con_internet_pct"] = round(_prop_w(net_flag, w)*100, 2)

    # Brechas en p.p. (edad, sexo, quintil)
    b_edad = tablas["Brecha por edad (exclusión)"]
    resumen["edad_peak"] = b_edad.loc[b_edad["Porcentaje"].idxmax()].to_dict()
    resumen["edad_valley"] = b_edad.loc[b_edad["Porcentaje"].idxmin()].to_dict()
    resumen["edad_brecha_pp"] = round(float(resumen["edad_peak"]["Porcentaje"] - resumen["edad_valley"]["Porcentaje"]), 2)

    if "Brecha por sexo (exclusión)" in tablas:
        t = tablas["Brecha por sexo (exclusión)"]
        pk = t.loc[t["Porcentaje"].idxmax()]; vl = t.loc[t["Porcentaje"].idxmin()]
        resumen["sexo_brecha_pp"] = round(float(pk["Porcentaje"] - vl["Porcentaje"]), 2)
        resumen["sexo_peak"] = pk.to_dict()

    if "Acceso a internet por quintil" in tablas:
        q = tablas["Acceso a internet por quintil"]
        resumen["quintil_brecha_pp"] = round(float(q.iloc[-1,1] - q.iloc[0,1]), 2)
        resumen["quintil_low"] = q.iloc[0].to_dict()
        resumen["quintil_high"] = q.iloc[-1].to_dict()

    return tablas, df, resumen

# =====================================
# Generación del informe Word extendido
# =====================================
def _p(doc: Document, text: str):
    if text:
        doc.add_paragraph(text)

def generar_informe_narrativo_tic(tablas: Dict[str, pd.DataFrame], anio: str = "2024", resumen: Optional[Dict[str, Any]] = None) -> BytesIO:
    doc = Document()
    doc.add_heading(f"Informe Analítico – Inclusión Digital TIC 4ºT {anio}", 0)
    _p(doc, "Este informe presenta un diagnóstico de inclusión/exclusión digital a partir del módulo TIC del cuarto trimestre, combinando información de hogares e individuos.")

    # 1. Definición
    doc.add_heading("1. Definición y enfoque", level=1)
    _p(doc, "Se operacionaliza la exclusión digital a partir de: (i) infraestructura del hogar (acceso a internet, disponibilidad de computadora) y (ii) uso individual de internet y computadora. Se construye, además, un índice ordinal (sin, parcial y total exclusión).")

    # 2. Resultados clave (con tablas en formato listado)
    doc.add_heading("2. Resultados clave", level=1)
    nombres_amigables = {
        "Uso de computadora": "Uso de computadora (individuos)",
        "Uso de celular": "Uso de celular (individuos)",
        "Uso de internet": "Uso de internet (individuos)",
        "Hogar con computadora": "Disponibilidad de computadora (hogares)",
        "Hogar con acceso a internet": "Acceso a internet (hogares)",
        "Exclusión – Nivel ordinal": "Exclusión digital (índice ordinal)",
        "Brecha por edad (exclusión)": "Brecha por edad (exclusión de personas)",
        "Brecha por sexo (exclusión)": "Brecha por sexo (exclusión de personas)",
        "Brecha por educación (exclusión)": "Brecha por nivel educativo (exclusión de personas)",
        "Acceso a internet por quintil": "Acceso a internet por quintil de ingreso del hogar",
        "Top 5 segmentos con mayor exclusión": "Top 5 segmentos con mayor exclusión (Edad × Educación × Sexo)",
    }

    for nombre, tabla in tablas.items():
        doc.add_heading(nombres_amigables.get(nombre, nombre), level=2)
        # formato viñetas
        for _, row in tabla.iterrows():
            vals = []
            for col in tabla.columns:
                vals.append(f"{col}: {row[col]}")
            _p(doc, "• " + " – ".join(vals))

    # 3. Conclusiones y análisis extendido
    doc.add_heading("3. Conclusiones y análisis extendido", level=1)
    if resumen:
        bullets = []

        # Panorama general
        if "exclusion_total_pct" in resumen:
            bullets.append(f"La exclusión digital total alcanza al {resumen['exclusion_total_pct']}% de las personas.")
        if "hogar_con_internet_pct" in resumen and "hogar_con_pc_pct" in resumen:
            bullets.append(f"En infraestructura, {resumen['hogar_con_internet_pct']}% de los hogares cuenta con internet y {resumen['hogar_con_pc_pct']}% dispone de computadora.")

        # Brechas
        if "edad_brecha_pp" in resumen and "edad_peak" in resumen and "edad_valley" in resumen:
            bullets.append(
                f"Por edad, la brecha entre el grupo más afectado ({resumen['edad_peak'].get('Edad','')}: {resumen['edad_peak'].get('Porcentaje','')}%) "
                f"y el menos afectado ({resumen['edad_valley'].get('Edad','')}: {resumen['edad_valley'].get('Porcentaje','')}%) es de "
                f"{resumen['edad_brecha_pp']} puntos porcentuales."
            )
        if "sexo_brecha_pp" in resumen:
            bullets.append(f"Existe una diferencia por sexo de {resumen['sexo_brecha_pp']} p.p. en la tasa de exclusión.")
        if "quintil_brecha_pp" in resumen and "quintil_low" in resumen and "quintil_high" in resumen:
            bullets.append(
                f"Por ingreso del hogar, el acceso a internet muestra una desigualdad de {resumen['quintil_brecha_pp']} p.p. entre "
                f"{resumen['quintil_low'].get('Quintil ingreso','Q1')} ({resumen['quintil_low'].get('% hogares con internet','')}%) y "
                f"{resumen['quintil_high'].get('Quintil ingreso','Q5')} ({resumen['quintil_high'].get('% hogares con internet','')}%)."
            )

        # Segmentos críticos
        if "Top 5 segmentos con mayor exclusión" in tablas:
            bullets.append("Los segmentos con mayor riesgo (combinando edad, educación y sexo) demandan intervenciones focalizadas en apropiación y equipamiento.")

        # Redacción final
        _p(doc, " ".join(bullets) if bullets else "Se observan brechas relevantes por edad, sexo, nivel educativo e ingreso.")

    _p(doc, "La evidencia sugiere que la exclusión digital no solo refleja rezagos tecnológicos, sino que también amplifica desigualdades educativas y de ingresos. Las políticas deben integrar conectividad de calidad, acceso a dispositivos y formación para usos significativos (educativos, laborales y administrativos).")

    # 4. Recomendaciones accionables
    doc.add_heading("4. Recomendaciones accionables", level=1)
    _p(doc, "• Priorizar expansión de conectividad fija en territorios con menor cobertura (barrios periféricos y localidades pequeñas).")
    _p(doc, "• Programas de acceso a dispositivos (computadoras) para hogares de Q1–Q2 con estudiantes en edad escolar/terciaria.")
    _p(doc, "• Alfabetización digital diferenciada por grupos etarios y trayectorias educativas, con tutorías presenciales y virtuales.")
    _p(doc, "• Articulación con instituciones educativas y de empleo para promover competencias digitales aplicadas (trámites, educación, trabajo remoto).")

    # 5. Limitaciones y próximas acciones
    doc.add_heading("5. Limitaciones y próximas acciones", level=1)
    _p(doc, "Los resultados dependen de la calidad de los registros y de la disponibilidad de variables según trimestre. Se sugiere complementar con análisis multivariados (p. ej., regresión logística) y comparativas interanuales para monitorear progresos 2017–2024.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

